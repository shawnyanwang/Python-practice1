__author__ = 'shawn.wang'

import wx
from docx import Document
import os


class MyFrame(wx.Frame):

    def __init__(self, parent):
        wx.Frame.__init__(self, parent, -1, 'Replace words in docx', size=(560, 380))
        self.folder_path = '~/'
        self.save_folder_path = '~/'
        panel = wx.Panel(self)
        sizer = wx.BoxSizer(wx.VERTICAL)
        panel.SetSizer(sizer)
        button_open = wx.Button(panel, -1, "Open", pos=(20, 18), size=(70, 35))
        self.Path_text = wx.StaticText(panel, -1, 'Path: '+self.folder_path, (100, 25), (400, 18), wx.ALIGN_LEFT)

        button_save_fold = wx.Button(panel, -1, "Save", pos=(20, 58), size=(70, 35))
        self.save_fold_text = wx.StaticText(panel, -1, 'Path: '+self.folder_path, (100, 67), (400, 48), wx.ALIGN_LEFT)

        self.label_search_content = wx.StaticText(panel, -1, 'Search content:', (20, 110), (400, 18), wx.ALIGN_LEFT)
        self.search_content = wx.TextCtrl(panel, -1, '', (20, 140), (500, 30), wx.ALIGN_LEFT)

        self.label_replace_content = wx.StaticText(panel, -1, 'Replace content:', (20, 180), (400, 18), wx.ALIGN_LEFT)
        self.replace_content = wx.TextCtrl(panel, -1, '', (20, 210), (500, 30), style=wx.TE_PROCESS_ENTER)

        button_run = wx.Button(panel, -1, "Replace", pos=(20, 255), size=(70, 35))

        self.Bind(wx.EVT_BUTTON, self.select_path, button_open)
        self.Bind(wx.EVT_BUTTON, self.select_save_fold, button_save_fold)
        self.Bind(wx.EVT_TEXT_ENTER, self.set_replace_content, self.replace_content)
        self.Bind(wx.EVT_BUTTON, self.batch_text_replace_doxc, button_run)
        self.Center()

    def select_path(self, evt):
        dialog = wx.DirDialog(None, "Select directory to open", self.folder_path, 0, (10, 10), wx.Size(400, 300))
        ret = dialog.ShowModal()
        self.folder_path = dialog.GetPath()+'/'
        if self.save_folder_path == '~/':
            self.save_folder_path = self.folder_path
            self.save_fold_text.LabelText = 'Path: '+self.save_folder_path
        self.Path_text.LabelText = 'Path: '+self.folder_path
        dialog.Destroy()

    def select_save_fold(self, evt):
        dialog = wx.DirDialog(None, "Select directory to open", self.folder_path, 0, (10, 10), wx.Size(400, 300))
        ret = dialog.ShowModal()
        self.save_folder_path = dialog.GetPath()+'/'
        self.save_fold_text.LabelText = 'Path: '+self.save_folder_path
        dialog.Destroy()

    def set_replace_content(self, evt):
        wx.MessageBox('Please confirm search and replace content', 'hint')

    def batch_text_replace_doxc(self, evt):
        path = self.folder_path
        original_content = self.search_content.GetValue()#"than current"
        replace_content = self.replace_content.GetValue()#"than reported"
        new_path = self.save_folder_path

        file_list = os.listdir(path)
        i = 0
        while i < len(file_list):
            item = file_list[i]
            if item.find('.docx') == -1:
                print 'delete: ', item
                file_list.remove(item)
            else:
                i += 1

        print '**************************************************'
        print '             Get the \'.docx\' files'
        print '**************************************************'
        for i in file_list:
            print i
        print '**************************************************'
        print '                   Start Replacing'
        print '**************************************************'

        for item in file_list:
            print "Processing: " + path+item
            document = Document(path+item)
            for paragraph in document.paragraphs:
                for run in paragraph.runs:
                    if original_content in run.text:
                        run.text = run.text.replace(original_content, replace_content)

            for table in document.tables:
                for cell in table._cells:
                    if original_content in cell.text:
                        for paragraph in cell.paragraphs:
                            if original_content in paragraph.text:
                                paragraph.text = paragraph.text.replace(original_content, replace_content)
                                paragraph.style.font.size = 130000
                                print paragraph.text
            document.save(new_path+item)

        print '**************************************************'
        print '                   Completed!'
        print '**************************************************'
        wx.MessageBox('Completed!', 'hint')

class MyApp(wx.App):

    def OnInit(self):
        self.frame = MyFrame(None)
        self.frame.Show(True)
        return True

app = MyApp()
app.MainLoop()

