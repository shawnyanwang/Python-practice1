__author__ = 'Shawn.wang'

from docx import Document
import os


def batch_text_replace_doxc(original_content="than current",
       insert_content = "than reported",
       path="C:\SW document\ETCO2 Test\ETCO2 Shawn(Yan)1.1/" + 'V1.2/',
       new_path="C:\SW document\ETCO2 Test\ETCO2 Shawn(Yan)1.1/" + 'V1.2/'):

    file_list = os.listdir(path)
    i = 0
    while i < len(file_list):
        item = file_list[i]
        if item.find('.docx') == -1:
            print 'delete: ', item
            file_list.remove(item)
        else:
            i += 1

    print '**************************************************'
    print '             Get the \'.docx\' files'
    print '**************************************************'
    for i in file_list:
        print i
    print '**************************************************'
    print '                   Start Replacing'
    print '**************************************************'

    for item in file_list:
        print "Processing: " + path+item
        document = Document(path+item)
        for paragraph in document.paragraphs:
            for run in paragraph.runs:
                if original_content in run.text:
                    run.text = run.text.replace(original_content, insert_content)

        for table in document.tables:
            for cell in table._cells:
                if original_content in cell.text:
                    cell
                for paragraph in cell.paragraphs:
                    if original_content in paragraph.text:
                        paragraph.text = paragraph.text.replace(original_content, insert_content)
                        paragraph.style.font.size = 130000
                        print paragraph.text
        document.save(new_path+item)

    print '**************************************************'
    print '                   Completed!'
    print '**************************************************'

if __name__ == "__main__":
    batch_text_replace_doxc()
